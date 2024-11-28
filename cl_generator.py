import os
import sqlite3
from dotenv import load_dotenv
from openai import OpenAI
from chatgpt_analysis import get_most_recent_pdf, extract_text_from_pdf  # Importing functions
from docx import Document  # Library for creating Word documents

# Load environment variables
load_dotenv()

# Initialize OpenAI client
client = OpenAI(api_key=os.environ['OPENAI_API_KEY'])

# Define directories
cv_folder = "./user_data/"
db_path = "./linkedin_jobs.db"
cover_letter_folder = "./cover_letters/"

# Ensure the cover letters folder exists
os.makedirs(cover_letter_folder, exist_ok=True)

def generate_cover_letter(job_id):
    """
    Generate a cover letter for a specific job ID and save it as a .docx file.
    """
    # Get the most recent CV
    cv_path = get_most_recent_pdf(cv_folder)
    print(f"Using CV file: {cv_path}")
    user_cv = extract_text_from_pdf(cv_path)

    # Connect to the SQLite database
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # Fetch the job description for the given job_id
    cursor.execute("""
    SELECT job_id, job_description
    FROM job_description
    WHERE job_id = ?
    """, (job_id,))
    job = cursor.fetchone()

    if not job:
        conn.close()
        raise ValueError(f"Job ID {job_id} not found in the job_description table.")

    job_id, job_description = job
    if not job_description:
        conn.close()
        raise ValueError(f"Job ID {job_id} has no job description.")

    try:
        # Generate the cover letter using ChatGPT
        print(f"Generating cover letter for job ID {job_id}...")
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert cover letter writer."},
                {
                    "role": "user",
                    "content": (
                        f"Using the following CV:\n{user_cv}\n\n"
                        f"Create a personalized cover letter for this job description:\n{job_description}"
                    )
                }
            ]
        )

        # Extract ChatGPT's generated cover letter
        cover_letter = response.choices[0].message.content.strip()

        # Save the cover letter to a Word document
        cover_letter_path = os.path.join(cover_letter_folder, f"{job_id}_cl.docx")
        document = Document()
        document.add_heading('Cover Letter', level=1)
        document.add_paragraph(cover_letter)
        document.save(cover_letter_path)

        print(f"Cover letter for job ID {job_id} saved to {cover_letter_path}")

    except Exception as e:
        print(f"Error generating cover letter for job ID {job_id}: {e}")

    finally:
        # Close the database connection
        conn.close()

